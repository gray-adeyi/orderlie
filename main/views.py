from django.shortcuts import render, redirect, get_object_or_404
from django.urls import reverse
from django.views.generic.edit import CreateView, FormView
from django.views.generic import DetailView, TemplateView, ListView
from django.contrib import messages
from django.views import View
import tempfile
from io import BytesIO
import logging
from openpyxl import Workbook
from docx import Document
from docx.shared import Inches
from django.http import HttpResponse
from . import models, forms
# Create your views here.

logger = logging.getLogger(__name__)


class Index(TemplateView):
    template_name = 'main/index.html'


class RegisterClass(CreateView):
    model = models.Class
    template_name = 'main/create_class.html'
    fields = ['name',
              'faculty',
              'department',
              'governor',
              'deputy_governor',
              'level',
              'session']

    def get_success_url(self):
        messages.success(self.request, 'New class successfully created')
        return reverse('main:class-info', kwargs={slug:self.model.slug})


class RegisterStudent(FormView):
    model = models.Class
    form_class = forms.StudentForm
    template_name = 'main/create_student.html'
    fields = '__all__'

    def form_valid(self, form):
        form.save()
        messages.success(self.request, 'Your data was successfully uploaded')
        slug = models.Class.objects.get(name=form.cleaned_data.get('student_class')).slug
        return redirect('main:class-info', slug=slug)


class ClassList(ListView):
    template_name = 'class_list.html'
    context_object_name = 'classes'
    model = models.Class


class ClassInfo(DetailView):
    model = models.Class
    template_name = 'main/class_info.html'

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['form'] = forms.StudentForm(initial={'student_class': self.model.objects.get(slug=self.kwargs['slug'])})
        return context


def download_class_data(request, slug):
    student_class = get_object_or_404(models.Class, slug=slug)
    members_list = list(student_class.students.all())
    members = [(
        idx + 1,
        member.last_name,
        member.first_name,
        member.middle_name,
        f"'{member.matric_no}'") for idx, member in enumerate(members_list)]
    members_name_merged = [(
        idx + 1,
        f'{member.last_name.upper()} {member.first_name.upper()} {member.middle_name.upper()}',
        f"'{member.matric_no}'") for idx, member in enumerate(members_list)]
    logger.debug(members)
    if request.GET.get("format") == "xlsx":
        wb = Workbook()
        sheet = wb.active
        sheet.title = 'Bio-Data'
        headers = ['S/N',
                   'Lastname',
                   'Firstname',
                   'Middlename',
                   'Matric No', ] # Headers if names are required to be discrete.
        headers_name_merged = ['S/N',
                   'Names',
                   'Matric No', ] # Headers if names are required to be merged.

        if request.GET.get("names") == "merged": # Logic to loop through `headers_name_merged` and `members_name_merged`
            # to populate the excel sheet.
            for index, row in enumerate(sheet.iter_rows(min_row=1, max_col=3, max_row=len(members_name_merged) + 1)):
                if index == 0:
                    for idx, cell in enumerate(row):
                        cell.value = headers_name_merged[idx]

            for row in range(2, len(members_name_merged) + 2):
                for idx, col in enumerate(headers_name_merged):
                    sheet.cell(row=row, column=idx+1, value=members_name_merged[row-2][idx])

        else: # alternate logic if names are required to be discrete.
            for index, row in enumerate(sheet.iter_rows(min_row=1, max_col=5, max_row=len(members) + 1)):
                if index == 0:
                    for idx, cell in enumerate(row):
                        cell.value = headers[idx]

            for row in range(2, student_class.students.all().count() + 2):
                for idx, col in enumerate(headers):
                    sheet.cell(row=row, column=idx+1, value=members[row-2][idx])

        # saving the file
        temp_file = BytesIO()
        wb.save(filename=temp_file)
        temp_file.seek(0)

        response = HttpResponse(
            content_type="application/vnd.ms-excel"
        )
        response["Content-Disposition"] = f"inline; filename={student_class.name}-bio-data.xlsx"
        response["Content-Transfer-Encoding"] = "binary"
        response.write(temp_file.read())
        temp_file.close()
        return response
    if request.GET.get("format") == "docx":
        doc = Document()
        doc.add_heading(f'{student_class.name} Bio-Data')
        p = doc.add_paragraph(
            f"Session: {student_class.get_session_display()}")
        p = doc.add_paragraph(f"Level: {student_class.get_level_display()}")
        p = doc.add_paragraph(f"Session: {student_class.governor}")
        p = doc.add_paragraph(f"Session: {student_class.deputy_governor}")

        if request.GET.get('names') == "merged":
            table = doc.add_table(rows=1, cols=3)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'S/N'
            hdr_cells[1].text = 'Names'
            hdr_cells[2].text = 'Matric No'
            for idx, member in enumerate(student_class.students.all()):
                idx += 1
                row_cells = table.add_row().cells
                row_cells[0].text = str(idx)
                row_cells[1].text = f"{member.last_name.upper()} {member.first_name.upper()} {member.middle_name.upper()}"
                row_cells[2].text = str(member.matric_no)
                logger.debug("I was executed")
        else:
            table = doc.add_table(rows=1, cols=5)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'S/N'
            hdr_cells[1].text = 'Lastname'
            hdr_cells[2].text = 'Firstname'
            hdr_cells[3].text = 'Middlename'
            hdr_cells[4].text = 'Matric No'
            for idx, member in enumerate(student_class.students.all()):
                idx += 1
                row_cells = table.add_row().cells
                row_cells[0].text = str(idx)
                row_cells[1].text = member.last_name.upper()
                row_cells[2].text = member.first_name.upper()
                row_cells[3].text = member.middle_name.upper()
                row_cells[4].text = str(member.matric_no)
        temp_file = BytesIO()
        doc.save(temp_file)
        temp_file.seek(0)

        response = HttpResponse(
            content_type="application/octet-stream"
        )
        response["Content-Disposition"] = f"inline; filename={student_class.name}-bio-data.docx"
        response["Content-Transfer-Encoding"] = "binary"
        response.write(temp_file.read())
        temp_file.close()
        return response
